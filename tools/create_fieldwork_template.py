from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "松阳寻踪_田野采集模板.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def write_cell(cell, text, bold=False, fill=None, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    if fill:
        set_cell_shading(cell, fill)


def add_table(doc, rows, widths, header=True):
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            write_cell(
                table.cell(row_index, column_index),
                value,
                bold=header and row_index == 0,
                fill=PALE_BLUE if header and row_index == 0 else None,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_prompt(doc, label, lines=2):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(label + "  " + "_" * (58 if lines == 1 else 30))
    set_font(run, size=10.5, bold=True, color=DARK_BLUE)
    for _ in range(lines - 1):
        blank = doc.add_paragraph("_" * 86)
        blank.paragraph_format.space_after = Pt(3)
        for item in blank.runs:
            set_font(item, size=10.5, color="666666")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run("巫通儿：松阳寻踪 | 田野采集模板")
    set_font(header_run, size=8.5, color="666666")
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("仅在取得对应授权后进入数字展览公开内容库")
    set_font(footer_run, size=8.5, color="666666")


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("巫通儿：松阳寻踪")
    set_font(run, size=26, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("传统村落数字田野展示 Demo | 田野采集与转化模板")
    set_font(run, size=13, color="555555")
    add_table(doc, [
        ["考察日期", "", "村落/路线", ""],
        ["记录人", "", "同行成员", ""],
        ["天气/时段", "", "设备与存储卡", ""],
        ["本日目标", "", "资料备份位置", ""],
    ], [1700, 2980, 1700, 2980], header=False)
    note = doc.add_table(rows=1, cols=1)
    set_table_widths(note, [9360])
    write_cell(note.cell(0, 0), "使用提醒：先取得知情同意，再录音、拍摄或采集精确位置。涉及个人身份、住址、敏感宗教仪式或未公开空间时，请标记为“仅内部研究”，不得直接进入公开展览。", fill=PALE_GOLD, size=10.5)
    doc.add_page_break()

    add_heading(doc, "1. 每日基础记录")
    add_table(doc, [
        ["字段", "填写内容"],
        ["今日路线", "起点 - 关键停留点 - 终点"],
        ["同行与接触对象", "姓名或代号；角色；是否同意公开"],
        ["观察重点", "建筑 / 水系 / 道路 / 山林 / 农事 / 仪式 / 日常交往"],
        ["今日最重要的三条发现", "1.\n2.\n3."],
        ["需要复访或核实的问题", ""],
    ], [2200, 7160])
    add_heading(doc, "2. 知情同意与使用范围", level=2)
    add_table(doc, [
        ["对象/资料编号", "口头同意", "书面同意", "可公开范围", "限制条件"],
        ["", "是 / 否", "是 / 否", "展览 / 课程 / 仅内部", "姓名、肖像、地点、原话、期限"],
        ["", "是 / 否", "是 / 否", "展览 / 课程 / 仅内部", ""],
        ["", "是 / 否", "是 / 否", "展览 / 课程 / 仅内部", ""],
    ], [1850, 1400, 1400, 2300, 2410])
    add_prompt(doc, "补充说明（无法取得同意、撤回、匿名化方式）", 2)

    doc.add_page_break()
    add_heading(doc, "3. 空间与节点记录卡")
    add_table(doc, [
        ["字段", "填写内容"],
        ["节点编号", "建议格式：SY-日期-序号"],
        ["节点名称 / 当地叫法", ""],
        ["类型", "民居 / 桥 / 古树 / 溪流 / 山路 / 晒场 / 祠堂 / 梯田 / 其他"],
        ["位置", "GPS 或相对位置；拍摄朝向；附近地标"],
        ["空间特征", "尺度、材质、颜色、地形关系、进入方式、停留行为"],
        ["人与空间的关系", "谁在使用；何时使用；与什么事件、劳动、仪式有关"],
        ["可转化为互动内容", "可探索动作、卡片主题、故事触发条件、环境音、关键词"],
        ["复核来源", "观察 / 访谈 / 史料 / 村民确认；资料编号"],
    ], [2200, 7160])
    add_heading(doc, "影像与测绘", level=2)
    add_table(doc, [
        ["文件编号", "内容与机位", "朝向/时间", "授权状态", "备注"],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
    ], [1600, 2900, 1650, 1650, 1560])

    doc.add_page_break()
    add_heading(doc, "4. 口述史与村落故事记录卡")
    add_table(doc, [
        ["字段", "填写内容"],
        ["受访者代号 / 身份", "例如：A01，村民；不在公开版记录姓名和住址"],
        ["访谈地点与时间", ""],
        ["话题", "古树、桥、迁徙、农事、节庆、建筑营造、山林守护等"],
        ["原话摘录", "请保留语境与语气；标记录音时间码"],
        ["故事摘要", "谁、在何处、发生了什么、为何重要"],
        ["可公开措辞", "用可确认、无夸张、可匿名的版本复述"],
        ["故事分类", "公共记忆 / 日常生活 / 生态经验 / 地方技艺 / 传说"],
        ["触发节点与关键词", "未来对应的 POI、卡片、角色台词或环境提示"],
    ], [2200, 7160])
    add_heading(doc, "敏感性与准确性检查", level=2)
    add_table(doc, [
        ["检查项", "结果"],
        ["是否含个人隐私、未公开地点或祭祀禁忌", "否 / 是，处理方式："],
        ["是否需要第二位来源核实", "否 / 是，待核问题："],
        ["原始记录文件", "录音 / 视频 / 照片 / 笔记编号："],
    ], [3500, 5860])

    doc.add_page_break()
    add_heading(doc, "5. 数字展览转化清单")
    add_table(doc, [
        ["素材/节点", "展览用途", "数据字段", "状态", "负责人"],
        ["", "3D 场景 / 研究卡 / 故事 / 环境音 / 交互动作", "标题、正文、关键词、来源、授权、文件路径", "待整理 / 待核实 / 可用", ""],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
    ], [1600, 2500, 2700, 1500, 1060])
    add_heading(doc, "最小可用资料包（每个公开节点）", level=2)
    for text in (
        "节点名称、当地称呼、坐标或相对位置、至少一张可用照片",
        "100 至 180 字经核实的说明文字，以及 3 至 5 个关键词",
        "至少一条可公开的故事或观察；原始来源及授权状态",
        "模型参考：正面、侧面、整体环境、近景材质各至少一张",
        "展览交互定义：靠近提示、触发按键、卡片主题、是否有环境音",
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        set_font(run, size=10.5)
    add_heading(doc, "6. 当日备份与复盘", level=1)
    add_table(doc, [
        ["项目", "确认"],
        ["原始照片、录音、视频已复制至两处存储", "是 / 否；位置："],
        ["文件已按日期-节点-序号命名", "是 / 否"],
        ["授权状态已写入资料清单", "是 / 否"],
        ["待核实问题和复访计划已同步", "是 / 否"],
    ], [4500, 4860])
    add_prompt(doc, "明日计划与风险提示", 3)
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
